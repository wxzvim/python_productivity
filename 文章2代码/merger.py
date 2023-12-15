from pathlib import Path, PurePath

from docx import Document


src_files = "./word样例文件"


p = Path(src_files)

files = [x for x in p.iterdir() if PurePath(x).match("*.docx")]

new_doc = Document()


def merger_files(doc_file: list):
    for file in sorted(doc_file):
        doc = Document(file)
        for word_page in doc.element.body:
            new_doc.element.body.append(word_page)

    new_doc.save(Path(src_files, "new.docx"))


def merger_without_format(doc_file: list):
    for file in sorted(doc_file):
        doc = Document(file)
        paras = doc.paragraphs
        for para in paras:
            new_para = new_doc.add_paragraph("")
            new_para.add_run(para.text)

    new_doc.save(Path(src_files, "news.docx"))


# merger_files(files)
merger_without_format(files)
