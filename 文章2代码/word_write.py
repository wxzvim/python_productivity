from docx import Document
from pathlib import Path

from docx.shared import RGBColor

word_files_path = "./word样例文件/"

p = Path(word_files_path)

mycontent = '''

Shiro 简介
照例又去官网扒了扒介绍：

Apache Shiro™ is a powerful and easy-to-use Java security framework that performs authentication, authorization, cryptography, and session management. With Shiro’s easy-to-understand API, you can quickly and easily secure any application – from the smallest mobile applications to the largest web and enterprise applications.
Apache Shiro™是一个强大且易用的Java安全框架,能够用于身份验证、授权、加密和会话管理。Shiro拥有易于理解的API,您可以快速、轻松地获得任何应用程序——从最小的移动应用程序到最大的网络和企业应用程序。
简而言之，Apache Shiro 是一个强大灵活的开源安全框架，可以完全处理身份验证、授权、加密和会话管理。

Shiro能到底能做些什么呢？

验证用户身份
用户访问权限控制，比如：1、判断用户是否分配了一定的安全角色。2、判断用户是否被授予完成某个操作的权限
在非 Web 或 EJB 容器的环境下可以任意使用Session API
可以响应认证、访问控制，或者 Session 生命周期中发生的事件
可将一个或以上用户安全数据源数据组合成一个复合的用户 “view”(视图)
支持单点登录(SSO)功能
支持提供“Remember Me”服务，获取用户关联信息而无需登录
···
为什么是 Shiro？
使用 Shiro 官方给了许多令人信服的原因，因为 Shiro 具有以下几个特点：

易于使用——易用性是项目的最终目标。应用程序安全非常令人困惑和沮丧,被认为是“不可避免的灾难”。如果你让它简化到新手都可以使用它,它就将不再是一种痛苦了。
全面——没有其他安全框架的宽度范围可以同Apache Shiro一样,它可以成为你的“一站式”为您的安全需求提供保障。
灵活——Apache Shiro可以在任何应用程序环境中工作。虽然在网络工作、EJB和IoC环境中可能并不需要它。但Shiro的授权也没有任何规范,甚至没有许多依赖关系。
Web支持——Apache Shiro拥有令人兴奋的web应用程序支持,允许您基于应用程序的url创建灵活的安全策略和网络协议(例如REST),同时还提供一组JSP库控制页面输出。
低耦合——Shiro干净的API和设计模式使它容易与许多其他框架和应用程序集成。你会看到Shiro无缝地集成Spring这样的框架, 以及Grails, Wicket, Tapestry, Mule, Apache Camel, Vaadin…等。
被广泛支持——Apache Shiro是Apache软件基金会的一部分。项目开发和用户组都有友好的网民愿意帮助。这样的商业公司如果需要Katasoft还提供专业的支持和服务。
有兴趣的可以去仔细看看官方的文档：【传送门】
Apache Shiro Features 特性
Apache Shiro是一个全面的、蕴含丰富功能的安全框架。下图为描述Shiro功能的框架图：

'''


def add_content_mode1(content):
    para = doc.add_paragraph().add_run(content)
    para.font.name = "仿宋"
    para.font.underline = True
    para.font.color.rgb = RGBColor(255, 128, 128)


doc = Document()
add_content_mode1(mycontent)
doc.save(Path(word_files_path, "new4.docx"))
